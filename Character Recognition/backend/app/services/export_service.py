"""
Export service — generates TXT, PDF, and DOCX from recognized text.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class ExportService:
    """Converts OCR output text to various file formats."""

    def to_txt(self, text: str) -> bytes:
        """Return UTF-8 encoded plain text."""
        return text.encode("utf-8")

    def to_pdf(self, text: str, title: str = "OCR Result") -> bytes:
        """Generate a PDF using reportlab."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4,
                                    leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            story = [
                Paragraph(f"<b>{title}</b>", styles["Title"]),
                Spacer(1, 0.5*cm),
            ]
            for line in text.split("\n"):
                story.append(Paragraph(line or "&nbsp;", styles["Normal"]))
                story.append(Spacer(1, 0.1*cm))

            doc.build(story)
            buf.seek(0)
            return buf.read()
        except ImportError:
            logger.warning("reportlab not installed; returning text as PDF placeholder.")
            return text.encode("utf-8")

    def to_docx(self, text: str, title: str = "OCR Result") -> bytes:
        """Generate a DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            doc.add_heading(title, level=1)
            for line in text.split("\n"):
                doc.add_paragraph(line)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()
        except ImportError:
            logger.warning("python-docx not installed; returning text bytes.")
            return text.encode("utf-8")

    def export(
        self,
        text: str,
        format: str = "txt",
        title: str = "OCR Result",
    ) -> tuple[bytes, str, str]:
        """
        Export text to the specified format.

        Returns:
            (content_bytes, media_type, file_extension)
        """
        fmt = format.lower()
        if fmt == "pdf":
            return self.to_pdf(text, title), "application/pdf", ".pdf"
        elif fmt == "docx":
            return self.to_docx(text, title), \
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document", \
                   ".docx"
        else:
            return self.to_txt(text), "text/plain; charset=utf-8", ".txt"
